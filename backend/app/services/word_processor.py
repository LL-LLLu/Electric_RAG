import os
import re
import json
import logging
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class WordProcessor:
    """Process Word documents to extract text chunks for RAG"""

    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
    TARGET_CHUNK_SIZE = 500  # Target tokens per chunk (approximately)
    CHUNK_OVERLAP = 50  # Overlap between chunks

    # Equipment tag pattern
    TAG_PATTERN = re.compile(r'\b[A-Z]{2,4}[-_\s]?[A-Z]?\d{1,4}[A-Z]?\b', re.IGNORECASE)

    def parse_file(self, file_path: str) -> Tuple[List[Dict], List[Dict]]:
        """Parse Word document and extract chunks

        Args:
            file_path: Path to the .docx file

        Returns:
            Tuple of (structured_data_list, text_chunks_list)
            Note: Word docs typically return empty structured_data
            unless they contain tables with equipment data
        """
        self._validate_file(file_path)

        try:
            doc = Document(file_path)
        except Exception as e:
            logger.error(f"Failed to open Word document {file_path}: {e}")
            raise ValueError(f"Unable to open Word document: {e}")

        structured_data = []
        chunks = []

        # Extract content by sections
        sections = self._extract_sections(doc)

        # Process each section into chunks
        for section in sections:
            section_chunks = self._create_section_chunks(section)
            chunks.extend(section_chunks)

        # Extract tables as separate structured data
        table_data = self._extract_tables(doc)
        structured_data.extend(table_data)

        # Update chunk indices
        for i, chunk in enumerate(chunks):
            chunk['chunk_index'] = i

        return structured_data, chunks

    def _validate_file(self, file_path: str) -> None:
        """Validate file exists and is within size limits"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_size = os.path.getsize(file_path)
        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(f"File too large: {file_size} bytes exceeds {self.MAX_FILE_SIZE} limit")

        if file_size == 0:
            raise ValueError(f"File is empty: {file_path}")

    def _extract_sections(self, doc: Document) -> List[Dict]:
        """Extract document content organized by section headings

        Returns list of sections with:
        - heading: The section heading text
        - heading_path: Full path like "3.0 RTU > 3.2 Alarms"
        - content: The section text content
        - level: Heading level (1, 2, 3, etc.)
        """
        sections = []
        current_section = {'heading': 'Document Start', 'heading_path': '', 'content': '', 'level': 0}
        heading_stack = []  # Track heading hierarchy

        for para in doc.paragraphs:
            # Check if this is a heading
            if para.style.name.startswith('Heading'):
                # Save previous section if it has content
                if current_section['content'].strip():
                    sections.append(current_section.copy())

                # Determine heading level
                try:
                    level = int(para.style.name.replace('Heading ', '').strip())
                except ValueError:
                    level = 1

                # Update heading stack
                while heading_stack and heading_stack[-1][1] >= level:
                    heading_stack.pop()
                heading_stack.append((para.text.strip(), level))

                # Build heading path
                heading_path = ' > '.join([h[0] for h in heading_stack])

                current_section = {
                    'heading': para.text.strip(),
                    'heading_path': heading_path,
                    'content': '',
                    'level': level
                }
            else:
                # Add content to current section
                if para.text.strip():
                    current_section['content'] += para.text.strip() + '\n'

        # Don't forget the last section
        if current_section['content'].strip():
            sections.append(current_section)

        return sections

    def _create_section_chunks(self, section: Dict) -> List[Dict]:
        """Split a section into appropriately sized chunks

        Args:
            section: Section dict with heading, heading_path, content

        Returns:
            List of chunk dicts ready for storage
        """
        chunks = []
        content = section['content']
        heading_path = section['heading_path']

        if not content:
            return chunks

        # Rough estimate: 1 token ≈ 4 characters
        chars_per_chunk = self.TARGET_CHUNK_SIZE * 4
        overlap_chars = self.CHUNK_OVERLAP * 4

        # Split content into sentences
        sentences = re.split(r'(?<=[.!?])\s+', content)

        current_chunk = []
        current_length = 0

        for sentence in sentences:
            sentence_length = len(sentence)

            if current_length + sentence_length > chars_per_chunk and current_chunk:
                # Create chunk from accumulated sentences
                chunk_content = ' '.join(current_chunk)
                equipment_tags = self._extract_equipment_tags(chunk_content)

                chunks.append({
                    'chunk_index': 0,  # Will be updated later
                    'content': f"Section: {heading_path}\n\n{chunk_content}" if heading_path else chunk_content,
                    'source_location': heading_path or section['heading'],
                    'equipment_tags': json.dumps(equipment_tags, default=str) if equipment_tags else None
                })

                # Start new chunk with overlap
                overlap_sentences = []
                overlap_length = 0
                for s in reversed(current_chunk):
                    if overlap_length + len(s) < overlap_chars:
                        overlap_sentences.insert(0, s)
                        overlap_length += len(s)
                    else:
                        break

                current_chunk = overlap_sentences
                current_length = overlap_length

            current_chunk.append(sentence)
            current_length += sentence_length

        # Don't forget the last chunk
        if current_chunk:
            chunk_content = ' '.join(current_chunk)
            equipment_tags = self._extract_equipment_tags(chunk_content)

            chunks.append({
                'chunk_index': 0,
                'content': f"Section: {heading_path}\n\n{chunk_content}" if heading_path else chunk_content,
                'source_location': heading_path or section['heading'],
                'equipment_tags': json.dumps(equipment_tags, default=str) if equipment_tags else None
            })

        return chunks

    def _extract_tables(self, doc: Document) -> List[Dict]:
        """Extract tables that might contain equipment data

        Returns list of structured data entries from tables
        """
        structured_data = []

        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) < 2:
                continue  # Skip tables without data rows

            # Get header row
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

            # Look for equipment column
            equipment_col = None
            for i, header in enumerate(headers):
                if any(pattern in header for pattern in ['equipment', 'tag', 'device', 'asset', 'item']):
                    equipment_col = i
                    break

            if equipment_col is None:
                continue  # No equipment column found

            # Extract data rows
            for row_idx, row in enumerate(table.rows[1:], start=1):
                cells = [cell.text.strip() for cell in row.cells]

                if len(cells) <= equipment_col:
                    continue

                equipment_tag = cells[equipment_col]
                if not equipment_tag:
                    continue

                # Build data dict from row
                row_data = {}
                for i, header in enumerate(headers):
                    if i < len(cells) and cells[i]:
                        row_data[header] = cells[i]

                structured_data.append({
                    'equipment_tag': equipment_tag.upper(),
                    'data_type': 'SCHEDULE_ENTRY',  # Default for table data
                    'data_json': json.dumps(row_data, default=str),
                    'source_location': f"Table {table_idx + 1}, Row {row_idx + 1}"
                })

        return structured_data

    def _extract_equipment_tags(self, text: str) -> List[str]:
        """Extract equipment tags from text using regex"""
        if not text:
            return []

        matches = self.TAG_PATTERN.findall(text)
        # Normalize and deduplicate
        tags = list(set(match.upper().replace(' ', '-').replace('_', '-') for match in matches))
        return sorted(tags)


# Singleton instance
word_processor = WordProcessor()
